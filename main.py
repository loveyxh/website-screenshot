import logging
from pathlib import Path
from typing import List, Dict, Optional
import pandas as pd
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
from selenium.common.exceptions import WebDriverException
from webdriver_manager.chrome import ChromeDriverManager
from docx import Document
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont
import os
import concurrent.futures
from math import ceil
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import threading

# 配置日志
log_file = 'website_screenshot.log'
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(log_file, encoding='utf-8'),
        logging.StreamHandler()
    ]
)

# 统计信息
class Stats:
    def __init__(self):
        self.total = 0
        self.success = 0
        self.failed = 0

    def add_result(self, success: bool):
        self.total += 1
        if success:
            self.success += 1
        else:
            self.failed += 1

    def __str__(self):
        return f"总计: {self.total}, 成功: {self.success}, 失败: {self.failed}"

class WebsiteScreenshotGenerator:
    def __init__(self, excel_path: str):
        """初始化截图生成器

        Args:
            excel_path: Excel文件路径
        """
        self.excel_path = Path(excel_path)
        self.screenshots_dir = Path('screenshots')
        self.screenshots_dir.mkdir(exist_ok=True)
        self.thread_local = threading.local()
        
    def get_driver(self):
        """为每个线程创建独立的WebDriver实例"""
        if not hasattr(self.thread_local, 'driver'):
            chrome_options = Options()
            chrome_options.add_argument('--headless')
            chrome_options.add_argument('--start-maximized')
            chrome_options.add_argument('--no-proxy-server')
            chrome_options.add_argument('--ignore-certificate-errors')
            
            self.thread_local.driver = webdriver.Chrome(
                service=Service(ChromeDriverManager().install()),
                options=chrome_options
            )
        return self.thread_local.driver

    def read_excel_data(self) -> List[Dict]:
        """读取Excel文件数据

        Returns:
            包含网站信息的字典列表
        """
        try:
            df = pd.read_excel(self.excel_path, sheet_name='sheet1')
            return df.to_dict('records')
        except Exception as e:
            logging.error(f'读取Excel文件失败: {e}')
            raise
    
    def generate_404_image(self, screenshot_path: Path):
        """生成404错误图片

        Args:
            screenshot_path: 保存图片的路径
        """
        # 创建一个800x600的灰色背景图片
        img = Image.new('RGB', (800, 600), color='#f0f0f0')
        draw = ImageDraw.Draw(img)
        
        # 添加404文字
        try:
            # 尝试使用系统字体
            font = ImageFont.truetype('arial.ttf', size=60)
        except:
            # 如果找不到字体，使用默认字体
            font = ImageFont.load_default()
            
        # 计算文字位置使其居中
        text = '404 Not Found'
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        
        x = (400 - text_width) // 2
        y = (400 - text_height) // 2
        
        # 绘制文字
        draw.text((x, y), text, font=font, fill='#333333')
        
        # 保存图片
        img.save(screenshot_path)
        logging.info(f'已生成404图片: {screenshot_path}')

    def take_screenshot(self, url: str, site_name: str, driver=None, max_retries: int = 3) -> Path:
        """对网站进行截图

        Args:
            url: 网站域名
            site_name: 网站名称
            driver: WebDriver实例
            max_retries: 最大重试次数

        Returns:
            截图文件路径
        """
        # 使用传入的driver或获取新的driver
        driver = driver or self.get_driver()
        
        # 补充协议头
        if not url.startswith(('http://', 'https://')):            
            url = url.strip()
            # 尝试先用https，如果失败再尝试http
            url = f'https://{url}'
        
        screenshot_path = self.screenshots_dir / f'{site_name}.png'
        
        for attempt in range(max_retries):
            try:
                driver.set_page_load_timeout(30)  # 设置页面加载超时时间
                driver.get(url)
                # 等待页面加载
                driver.implicitly_wait(10)
                # 获取当前URL，检查是否发生跳转
                current_url = driver.current_url
                if current_url != url:
                    logging.info(f'网站发生跳转: {url} -> {current_url}')
                    url = current_url
                # 设置固定的窗口大小为800x600
                driver.set_window_size(800, 600)
                # 截图
                driver.save_screenshot(str(screenshot_path))
                logging.info(f'成功截图: {url}')
                return screenshot_path
            except WebDriverException as e:
                if attempt < max_retries - 1:
                    logging.warning(f'第{attempt + 1}次截图失败 {url}: {e}, 准备重试...')
                    # 如果是https失败，尝试http
                    if url.startswith('https://') and attempt == 0:
                        url = f'http://{url[8:]}'
                    continue
                else:
                    logging.error(f'截图失败 {url}: {e}, 已达到最大重试次数，生成404图片')
                    self.generate_404_image(screenshot_path)
                    return screenshot_path
    
    def process_website(self, item: Dict) -> tuple:
        """处理单个网站的截图，用于并发执行"""
        driver = self.get_driver()
        try:
            logging.info(f"开始处理序号 {item['序号']} 的网站: {item['网站名称']}")
            screenshot_path = self.take_screenshot(
                item['网站域名'], 
                item['网站名称'],
                driver=driver
            )
            return (item, True, screenshot_path)
        except Exception as e:
            logging.error(f"处理序号 {item['序号']} 的网站失败 {item['网站名称']}: {e}")
            # 生成404图片
            screenshot_path = self.screenshots_dir / f'{item["网站名称"]}.png'
            self.generate_404_image(screenshot_path)
            return (item, False, screenshot_path)

    def generate_word_document(self, data: List[tuple], output_path: str):
        """生成优化格式的Word文档"""
        doc = Document()
        
        for item, success, screenshot_path in data:
            # 添加标题和内容
            for field in ['序号', '网站名称', '网站域名']:
                p = doc.add_paragraph()
                # 标题加粗
                run = p.add_run(f'{field}: ')
                run.bold = True
                run.font.size = Pt(14)  # 增大字体大小
                # 内容不加粗
                run = p.add_run(f'{item[field]}')
                run.bold = False
                run.font.size = Pt(14)  # 保持字体大小一致
            
            # 添加截图
            if success and screenshot_path and Path(screenshot_path).exists():
                doc.add_picture(str(screenshot_path), width=Inches(6))
            
            # 添加占满一行的分隔线
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            p.add_run('_' * 80)  # 缩短分隔线长度
        
        # 保存文档
        doc.save(output_path)
        logging.info(f'文档已保存: {output_path}')

    def process(self, output_base_name: str):
        """使用并发处理所有网站并生成分页文档"""
        try:
            # 读取数据
            data = self.read_excel_data()
            stats = Stats()
            
            # 创建文档缓存字典，用于存储不同页的文档对象
            doc_cache = {}
            page_size = 100
            
            # 使用线程池并发处理网站
            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
                future_to_item = {
                    executor.submit(self.process_website, item): item 
                    for item in data
                }
                
                # 创建一个锁用于同步文档写入操作
                doc_lock = threading.Lock()
                
                for idx, future in enumerate(concurrent.futures.as_completed(future_to_item)):
                    item, success, screenshot_path = future.result()
                    stats.add_result(success)
                    
                    # 计算当前记录应该写入哪个文档
                    page_num = idx // page_size
                    start_idx = page_num * page_size
                    end_idx = start_idx + page_size
                    output_path = f"{output_base_name}({start_idx}-{end_idx}).docx"
                    
                    # 使用锁保护文档操作
                    with doc_lock:
                        # 如果文档不存在，创建新文档
                        if page_num not in doc_cache:
                            # 如果文档已存在，则加载它
                            if os.path.exists(output_path):
                                doc_cache[page_num] = Document(output_path)
                            else:
                                doc_cache[page_num] = Document()
                        
                        # 将当前记录写入对应的文档
                        doc = doc_cache[page_num]
                        for field in ['序号', '网站名称', '网站域名']:
                            p = doc.add_paragraph()
                            run = p.add_run(f'{field}: {item[field]}')
                            run.bold = True
                            run.font.size = Pt(14)
                        
                        # 添加截图
                        if success and screenshot_path and Path(screenshot_path).exists():
                            doc.add_picture(str(screenshot_path), width=Inches(6))
                        
                        # 添加分隔线
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        p.add_run('_' * 120)
                        
                        # 立即保存文档
                        doc.save(output_path)
                        logging.info(f'已更新文档: {output_path}')
                        
                        # 如果当前页已满，从缓存中移除文档
                        if (idx + 1) % page_size == 0:
                            del doc_cache[page_num]
            
            logging.info(f'处理完成，统计信息: {stats}')
            
        finally:
            # 关闭所有浏览器实例
            if hasattr(self.thread_local, 'driver'):
                self.thread_local.driver.quit()

def main():
    try:
        os.environ['no_proxy'] = '*'
        generator = WebsiteScreenshotGenerator('list.xlsx')
        generator.process('网站截图报告')
        logging.info('处理完成')
    except Exception as e:
        logging.error(f'程序执行失败: {e}')

if __name__ == '__main__':
    main()